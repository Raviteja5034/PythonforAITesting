"""
rag_pipeline.py — RAG (Retrieval-Augmented Generation) core logic
Supports both Groq (cloud, fast, free) and Ollama (local).
"""

import os
import re
import sys
from pathlib import Path

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.prompts import PromptTemplate

sys.path.insert(0, os.path.dirname(__file__))
from config import (
    VECTORSTORE_PATH, COLLECTION_NAME, EMBEDDING_MODEL,
    OLLAMA_BASE_URL, LLM_MODEL, TOP_K_DOCS,
    USE_GROQ, GROQ_API_KEY, GROQ_MODEL, MANIFEST_PATH,
)
from manifest import load_manifest

# ── Prompt Template ────────────────────────────────────────────────────────────
# This tells the LLM exactly how to behave: answer ONLY from the documents,
# say "I don't know" if the answer isn't there — avoids hallucinations
RAG_PROMPT_TEMPLATE = """You are the BPI Q&A assistant helping SIT and UAT testers understand system behavior, user stories, and design documents.

Use only the document context below to answer the user's question.
The context may contain user stories written in GIVEN/WHEN/THEN format — extract and explain them clearly.
If the context contains the answer, answer clearly and concisely from that context.
If the context is empty, weak, or unrelated to the question, reply exactly with: I could not find this information in the loaded documents.
Do not use outside knowledge.
Do not guess.
Do not invent document-specific facts that are not present in the context.
Be concise, factual, and helpful.

Context from documents:
-----------------------
{context}
-----------------------

Tester's Question: {question}

Answer:"""

RAG_PROMPT = PromptTemplate(
    template=RAG_PROMPT_TEMPLATE,
    input_variables=["context", "question"],
)

# ── Singleton instances (loaded once, reset on re-ingest) ─────────────────────
_embeddings  = None
_vectorstore = None
_llm_instance = None
_chain_lock  = __import__("threading").Lock()  # thread-safe reset


def _get_embeddings():
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
    return _embeddings


def _get_vectorstore():
    global _vectorstore
    if _vectorstore is None:
        if not os.path.exists(VECTORSTORE_PATH):
            raise FileNotFoundError(
                f"Vector store not found at '{VECTORSTORE_PATH}'. "
                "Run 'python ingest.py' first to load your documents."
            )
        _vectorstore = Chroma(
            collection_name=COLLECTION_NAME,
            persist_directory=VECTORSTORE_PATH,
            embedding_function=_get_embeddings(),
        )
    return _vectorstore


def reset_pipeline():
    """
    Fully reset the RAG pipeline after re-ingestion.
    Called by watcher after vectorstore is rebuilt on disk.
    Forces fresh Chroma connection + new retriever + new chain on next question.
    """
    global _vectorstore, _llm_instance
    with _chain_lock:
        # Close existing vectorstore connection if possible
        try:
            if _vectorstore is not None:
                _vectorstore._client.reset()
        except Exception:
            pass
        _vectorstore = None
        _llm_instance = None
    print("[RAG] Pipeline reset — will reload fresh vectorstore on next question.")


def _build_llm():
    """Return Groq or Ollama LLM based on config."""
    if USE_GROQ:
        if not GROQ_API_KEY:
            raise ValueError("GROQ_API_KEY is not set. Configure it as an environment variable before starting the app.")
        from langchain_groq import ChatGroq
        print(f"[RAG] Using Groq API — model: {GROQ_MODEL}")
        return ChatGroq(
            api_key=GROQ_API_KEY,
            model=GROQ_MODEL,
            temperature=0.1,
            max_tokens=768,
        )
    else:
        from langchain_ollama import OllamaLLM
        print(f"[RAG] Using Ollama local — model: {LLM_MODEL}")
        return OllamaLLM(
            model=LLM_MODEL,
            base_url=OLLAMA_BASE_URL,
            temperature=0.1,
            num_predict=768,
            num_ctx=4096,
        )


def _normalize_question(question: str) -> str:
    """Normalize common user spellings/aliases before retrieval."""
    normalized = question.strip()
    replacements = {
        "grok": "groq",
        "grok key": "groq api key",
        "grok api": "groq api",
        "pavankalyan": "pawan kalyan",
    }
    lower_normalized = normalized.lower()
    for wrong, right in replacements.items():
        lower_normalized = lower_normalized.replace(wrong, right)
    return lower_normalized


def _is_valid_doc(doc) -> bool:
    content = getattr(doc, "page_content", None)
    return isinstance(content, str) and bool(content.strip())


def _safe_document(doc):
    """Return a sanitized LangChain Document or None."""
    try:
        from langchain.schema import Document as LCDocument
        content = getattr(doc, "page_content", None)
        if not isinstance(content, str):
            return None
        content = content.strip()
        if not content:
            return None
        metadata = getattr(doc, "metadata", {}) or {}
        return LCDocument(page_content=content, metadata=dict(metadata))
    except Exception:
        return None


def _sanitize_documents(docs):
    safe_docs = []
    for doc in docs or []:
        safe_doc = _safe_document(doc)
        if safe_doc is not None:
            safe_docs.append(safe_doc)
    return safe_docs


def _keyword_fallback_documents(question: str):
    """Fallback lexical search over currently indexed source files for exact and near matches."""
    normalized = _normalize_question(question)
    tokens = [token for token in re.findall(r"[a-zA-Z0-9]+", normalized) if len(token) > 2]
    if not tokens:
        return []

    manifest = load_manifest(MANIFEST_PATH)
    ranked_docs = []
    docs_root = Path(os.path.abspath(os.environ.get("BPI_DOCS_PATH", ""))) if os.environ.get("BPI_DOCS_PATH") else None

    def token_hit_score(token: str, text: str) -> int:
        if token in text:
            return 3
        for word in text.split():
            if token in word or word in token:
                return 2
            if abs(len(word) - len(token)) <= 2:
                mismatches = sum(1 for a, b in zip(word, token) if a != b)
                mismatches += abs(len(word) - len(token))
                if mismatches <= 2:
                    return 1
        return 0

    for rel_path, entry in manifest.items():
        if entry.get("status") != "indexed":
            continue

        file_path = docs_root / rel_path if docs_root else Path(rel_path)
        if not file_path.exists():
            continue

        suffix = file_path.suffix.lower()
        if suffix not in {".txt", ".md", ".docx"}:
            continue

        try:
            if suffix == ".docx":
                from docx import Document
                text = "\n".join(p.text for p in Document(str(file_path)).paragraphs if p.text.strip())
            else:
                text = file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            continue

        normalized_text = re.sub(r"\s+", " ", text.lower())
        score = sum(token_hit_score(token, normalized_text) for token in tokens)
        freshness_bonus = 1 if entry.get("last_attempted_at") == entry.get("ingested_at") else 0
        score += freshness_bonus
        if score <= 0:
            continue

        from langchain.schema import Document as LCDocument
        ranked_docs.append((
            score,
            LCDocument(
                page_content=text[:4000],
                metadata={
                    "source_file": entry.get("source_file", rel_path),
                    "relative_path": rel_path,
                    "fallback_keyword_match": True,
                    "keyword_score": score,
                },
            )
        ))

    ranked_docs.sort(key=lambda item: item[0], reverse=True)
    return [doc for _, doc in ranked_docs[:TOP_K_DOCS]]


def _retrieve_documents(question: str):
    """Retrieve documents with keyword-first priority, then semantic search."""
    vectorstore = _get_vectorstore()
    normalized_question = _normalize_question(question)
    keyword_docs = _keyword_fallback_documents(normalized_question)

    if keyword_docs:
        primary_retriever = vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": TOP_K_DOCS},
        )
        semantic_docs = primary_retriever.invoke(normalized_question)
        existing_files = {doc.metadata.get("source_file") for doc in keyword_docs}
        for sdoc in semantic_docs:
            if sdoc.metadata.get("source_file") not in existing_files:
                keyword_docs.append(sdoc)
        return _sanitize_documents(keyword_docs[:TOP_K_DOCS])

    primary_retriever = vectorstore.as_retriever(
        search_type="similarity",
        search_kwargs={"k": TOP_K_DOCS},
    )
    docs = _sanitize_documents(primary_retriever.invoke(normalized_question))
    if docs:
        return docs[:TOP_K_DOCS]

    fallback_retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={"k": 6, "fetch_k": 12},
    )
    return _sanitize_documents(fallback_retriever.invoke(normalized_question))


def get_llm():
    global _llm_instance
    with _chain_lock:
        if _llm_instance is None:
            _llm_instance = _build_llm()
    return _llm_instance


def ask_question(question: str) -> dict:
    """
    Ask a question and return the answer + source documents.
    Returns: { "answer": str, "sources": [{"file": str, "snippet": str}] }
    """
    llm = get_llm()
    normalized_question = _normalize_question(question)
    docs = _sanitize_documents(_retrieve_documents(question))
    if not docs:
        return {"answer": "I could not find this information in the loaded documents.", "sources": []}

    context = "\n\n".join(doc.page_content for doc in docs if _is_valid_doc(doc))
    if not context.strip():
        return {"answer": "I could not find this information in the loaded documents.", "sources": []}

    prompt = RAG_PROMPT.format(context=context, question=normalized_question)
    response = llm.invoke(prompt)
    answer = response.content.strip() if hasattr(response, "content") else str(response).strip()

    # Collect unique source files and a short snippet from each
    sources = []
    seen_files = set()
    for doc in docs:
        fname = doc.metadata.get("source_file", doc.metadata.get("source", "Unknown"))
        if fname not in seen_files:
            seen_files.add(fname)
            snippet = doc.page_content[:200].replace("\n", " ").strip()
            sources.append({"file": fname, "snippet": snippet + "…"})

    return {"answer": answer, "sources": sources}


def get_vectorstore_count() -> int:
    try:
        vs = _get_vectorstore()
        return vs._collection.count()
    except Exception:
        return 0


def is_vectorstore_ready() -> bool:
    """Check whether documents have been ingested and roughly match manifest state."""
    try:
        count = get_vectorstore_count()
        if count <= 0:
            return False
        manifest = load_manifest(MANIFEST_PATH)
        expected = sum(item.get("chunk_count", 0) for item in manifest.values())
        if expected > 0:
            drift = abs(expected - count) / expected
            if drift > 0.2:
                print(f"[RAG][WARN] Manifest expects {expected} chunks but collection has {count}")
        return True
    except Exception:
        return False


def prewarm():
    """
    Pre-load the model at server startup so the first real question is fast.
    Sends a trivial prompt to Ollama to pull the model into RAM.
    """
    try:
        print(f"  Pre-warming model '{LLM_MODEL}' — please wait...")
        get_qa_chain()
        # Minimal warm-up: initialises embeddings + retriever at startup
        print(f"  Model chain ready.")
    except Exception as e:
        print(f"  [WARN] Pre-warm failed: {e}")
