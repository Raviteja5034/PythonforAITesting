"""
ingest.py — Document ingestion pipeline
Run this script once (or whenever you add new docs) to:
  1. Load supported files from DOCS_PATH
  2. Split them into chunks
  3. Embed with HuggingFace
  4. Synchronise ChromaDB with latest document set
"""

import os
import sys
import time
import argparse
from pathlib import Path

from langchain_community.document_loaders import (
    PyPDFLoader,
    Docx2txtLoader,
    TextLoader,
    UnstructuredPowerPointLoader,
    UnstructuredExcelLoader,
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "backend"))
sys.path.insert(0, os.path.dirname(__file__))
from config import (
    DOCS_PATH, VECTORSTORE_PATH, COLLECTION_NAME,
    EMBEDDING_MODEL, CHUNK_SIZE, CHUNK_OVERLAP, EXCLUDED_FILES, MANIFEST_PATH,
)
from manifest import load_manifest, save_manifest, changed_files

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".pptx", ".xlsx", ".txt", ".md"}
RETRY_ATTEMPTS = 3
RETRY_SLEEP_SECONDS = 2


class IngestStats:
    def __init__(self):
        self.loaded_files = 0
        self.failed_files = 0
        self.skipped_files = 0
        self.deleted_files = 0
        self.chunk_count = 0
        self.errors = []
        self.sync_status = "updated"


def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def get_vectorstore(embeddings):
    return Chroma(
        collection_name=COLLECTION_NAME,
        persist_directory=VECTORSTORE_PATH,
        embedding_function=embeddings,
    )


def _enrich_metadata(docs, file_info: dict):
    valid_docs = []
    for doc in docs:
        content = getattr(doc, "page_content", None)
        if not isinstance(content, str) or not content.strip():
            continue
        doc.page_content = content.strip()
        doc.metadata["source_file"] = file_info["source_file"]
        doc.metadata["relative_path"] = file_info["relative_path"]
        doc.metadata["sha256"] = file_info["sha256"]
        doc.metadata["source_mtime"] = file_info["mtime"]
        valid_docs.append(doc)
    return valid_docs


def _load_docx_as_paragraph_documents(path: str, file_info: dict):
    from langchain.schema import Document
    from docx import Document as WordDocument

    docx = WordDocument(path)
    paragraphs = [p.text.strip() for p in docx.paragraphs if p.text and p.text.strip()]
    if not paragraphs:
        raise ValueError("DOCX file produced no readable paragraphs")

    docs = []
    buffer = []
    char_count = 0
    target_size = min(CHUNK_SIZE, 350)
    for paragraph in paragraphs:
        if buffer and (char_count + len(paragraph) + 2) > target_size:
            docs.append(Document(page_content="\n\n".join(buffer), metadata={}))
            buffer = []
            char_count = 0
        buffer.append(paragraph)
        char_count += len(paragraph) + 2
    if buffer:
        docs.append(Document(page_content="\n\n".join(buffer), metadata={}))

    return _enrich_metadata(docs, file_info)


def _load_single_file(file_info: dict):
    path = file_info["absolute_path"]
    ext = Path(path).suffix.lower()
    loader = None
    if ext == ".pdf":
        loader = PyPDFLoader(path)
    elif ext == ".pptx":
        loader = UnstructuredPowerPointLoader(path)
    elif ext == ".xlsx":
        loader = UnstructuredExcelLoader(path)
    elif ext == ".docx":
        loader = None
    else:
        try:
            loader = TextLoader(path, encoding="utf-8")
            loader.load()
            loader = TextLoader(path, encoding="utf-8")
        except Exception:
            loader = TextLoader(path, encoding="cp1252")

    last_error = None
    for attempt in range(1, RETRY_ATTEMPTS + 1):
        try:
            if ext == ".docx":
                loaded = _load_docx_as_paragraph_documents(path, file_info)
            else:
                loaded = loader.load()
                if not loaded:
                    raise ValueError("File produced no readable content")
                loaded = _enrich_metadata(loaded, file_info)
            if not loaded:
                raise ValueError("File produced no readable content")
            return loaded, None
        except Exception as exc:
            last_error = exc
            if attempt < RETRY_ATTEMPTS:
                time.sleep(RETRY_SLEEP_SECONDS * attempt)
    return [], last_error


def split_documents(docs):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=min(CHUNK_SIZE, 350),
        chunk_overlap=min(CHUNK_OVERLAP, 50),
        separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""],
    )
    return splitter.split_documents(docs)


def synchronize_documents(docs_path: str):
    path = Path(docs_path)
    if not path.exists():
        print(f"[ERROR] Path does not exist: {docs_path}")
        sys.exit(1)

    manifest = load_manifest(MANIFEST_PATH)
    changes = changed_files(manifest, docs_path, SUPPORTED_EXTENSIONS, EXCLUDED_FILES)
    stats = IngestStats()

    print("Change summary:")
    print(f"  Added    : {len(changes['added'])}")
    print(f"  Modified : {len(changes['modified'])}")
    print(f"  Deleted  : {len(changes['deleted'])}")
    print(f"  Unchanged: {len(changes['unchanged'])}")

    if not changes["added"] and not changes["modified"] and not changes["deleted"]:
        print("Nothing to update — vector store already matches the latest documents.")
        stats.sync_status = "noop"
        return stats

    embeddings = get_embeddings()
    vectorstore = get_vectorstore(embeddings)
    collection = vectorstore._collection

    for deleted in changes["deleted"]:
        try:
            collection.delete(where={"relative_path": deleted["relative_path"]})
            manifest.pop(deleted["relative_path"], None)
            stats.deleted_files += 1
            print(f"  [DELETE] Removed from index: {deleted['relative_path']}")
        except Exception as exc:
            manifest[deleted["relative_path"]] = {
                **deleted,
                "status": "delete_failed",
                "last_error": str(exc),
                "last_attempted_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            stats.errors.append(f"Delete failed for {deleted['relative_path']}: {exc}")
            print(f"  [WARN] Could not delete {deleted['relative_path']}: {exc}")

    files_to_process = changes["added"] + changes["modified"]
    for file_info in files_to_process:
        loaded, error = _load_single_file(file_info)
        if error is not None:
            manifest[file_info["relative_path"]] = {
                "relative_path": file_info["relative_path"],
                "source_file": file_info["source_file"],
                "sha256": file_info["sha256"],
                "mtime": file_info["mtime"],
                "size": file_info["size"],
                "chunk_count": manifest.get(file_info["relative_path"], {}).get("chunk_count", 0),
                "ingested_at": manifest.get(file_info["relative_path"], {}).get("ingested_at"),
                "status": "load_failed",
                "last_error": str(error),
                "last_attempted_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            stats.failed_files += 1
            stats.errors.append(f"Load failed for {file_info['relative_path']}: {error}")
            print(f"  [WARN] Could not load {file_info['source_file']}: {error}")
            continue

        chunks = split_documents(loaded)
        if not chunks:
            manifest[file_info["relative_path"]] = {
                "relative_path": file_info["relative_path"],
                "source_file": file_info["source_file"],
                "sha256": file_info["sha256"],
                "mtime": file_info["mtime"],
                "size": file_info["size"],
                "chunk_count": 0,
                "ingested_at": manifest.get(file_info["relative_path"], {}).get("ingested_at"),
                "status": "empty",
                "last_error": "No usable content extracted",
                "last_attempted_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            stats.skipped_files += 1
            stats.errors.append(f"No chunks generated for {file_info['relative_path']}")
            print(f"  [WARN] No usable content in {file_info['source_file']}")
            continue

        try:
            collection.delete(where={"relative_path": file_info["relative_path"]})
        except Exception:
            pass

        ids = []
        for index, chunk in enumerate(chunks):
            doc_id = f"{file_info['sha256'][:12]}_{index:04d}"
            chunk.metadata["doc_id"] = file_info["sha256"][:12]
            chunk.metadata["chunk_index"] = index
            ids.append(doc_id)

        try:
            vectorstore.add_documents(chunks, ids=ids)
            manifest[file_info["relative_path"]] = {
                "relative_path": file_info["relative_path"],
                "source_file": file_info["source_file"],
                "sha256": file_info["sha256"],
                "mtime": file_info["mtime"],
                "size": file_info["size"],
                "chunk_count": len(chunks),
                "ingested_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                "status": "indexed",
                "last_error": None,
                "last_attempted_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            stats.loaded_files += 1
            stats.chunk_count += len(chunks)
            print(f"  [OK] Indexed: {file_info['source_file']} ({len(chunks)} chunk(s))")
        except Exception as exc:
            manifest[file_info["relative_path"]] = {
                "relative_path": file_info["relative_path"],
                "source_file": file_info["source_file"],
                "sha256": file_info["sha256"],
                "mtime": file_info["mtime"],
                "size": file_info["size"],
                "chunk_count": manifest.get(file_info["relative_path"], {}).get("chunk_count", 0),
                "ingested_at": manifest.get(file_info["relative_path"], {}).get("ingested_at"),
                "status": "index_failed",
                "last_error": str(exc),
                "last_attempted_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
            stats.failed_files += 1
            stats.errors.append(f"Index failed for {file_info['relative_path']}: {exc}")
            print(f"  [WARN] Could not index {file_info['source_file']}: {exc}")

    save_manifest(manifest, MANIFEST_PATH)
    print(f"Manifest saved: {MANIFEST_PATH}")
    print(f"Collection count: {collection.count()}")
    return stats


def main():
    parser = argparse.ArgumentParser(description="BPI Q&A — Document Ingestion")
    parser.add_argument("--path", default=DOCS_PATH, help="Path to documents folder or network share")
    args = parser.parse_args()

    print("=" * 60)
    print("  BPI Q&A — Document Ingestion Pipeline")
    print("=" * 60)
    print(f"Source path: {args.path}\n")

    stats = synchronize_documents(args.path)
    print("\nSummary:")
    print(f"  Indexed files : {stats.loaded_files}")
    print(f"  Deleted files : {stats.deleted_files}")
    print(f"  Failed files  : {stats.failed_files}")
    print(f"  Skipped files : {stats.skipped_files}")
    print(f"  Chunks added  : {stats.chunk_count}")
    if stats.errors:
        print("  Errors:")
        for err in stats.errors:
            print(f"    - {err}")

    print("\nIngestion complete! You can now start the backend server.")
    print("  python backend/main.py")


if __name__ == "__main__":
    main()
